from docx import Document
from docx.shared import Inches

# Initialize the Word document
doc = Document()

# Define section titles
sections = {
    "Login and Logout": [],
    "Login and Buy 4 Items": [],
    "Multiple Scenarios - 1 Feature": []
}

# Assign screenshot filenames to each section (Replace these with actual filenames)
login_logout_screenshots = ["screenshots/login_button.png",
                            "screenshots/login_page.png",
                            "screenshots/login_successful.png",
                            "screenshots/logout_button.png",
                            "screenshots/logout_page.png",
                            "screenshots/logout_successful.png",
                            "screenshots/password.png",
                            "screenshots/username.png"]
login_buy_4_items_screenshots = ["screenshots/card_details.png",
                                 "screenshots/cart.png",
                                 "screenshots/checkout.png",
                                 "screenshots/address_empty.png",
                                 "screenshots/address_added.png",
                                 "screenshots/payment_added.png",
                                 "screenshots/payment_empty.png",
                                 "screenshots/place_order.png",
                                 "screenshots/products.png",
                                 "screenshots/products_0.png",
                                 "screenshots/products_3.png",
                                 "screenshots/products_2.png",
                                 "screenshots/products_1.png",
                                 "screenshots/review_order.png"]
multiple_scenarios_screenshots = ["screenshots/invalid_card.png",
                                  "screenshots/before_delete.png",
                                  "screenshots/after_deletion.png",
                                  "screenshots/card_details.png"]

# Map screenshots to sections
sections["Login and Logout"] = login_logout_screenshots
sections["Login and Buy 4 Items"] = login_buy_4_items_screenshots
sections["Multiple Scenarios - 1 Feature"] = multiple_scenarios_screenshots

# Generate document content
for title, screenshots in sections.items():
    doc.add_heading(title, level=1)  # Add section title
    if title == "Login and Logout":
        doc.add_paragraph("La funcionalidad del feature permite a un usuario iniciar sesión en la aplicación. En el escenario descrito, el usuario accede a la pantalla de inicio de sesión, ingresa sus credenciales (correo y contraseña) y presiona el botón de inicio de sesión. Si las credenciales son correctas, el sistema muestra la pantalla de inicio. Finalmente, el usuario cierra sesión (logout).")
    elif title == "Login and Buy 4 Items":
        doc.add_paragraph("La funcionalidad del feature permite a un usuario autenticado realizar una compra en la aplicación. En el escenario descrito, el usuario inicia sesión, selecciona 4 productos y los agrega al carrito. Luego, procede al pago, ingresa una dirección de envío y un método de pago. Antes de finalizar, revisa la orden y la confirma. Finalmente, el sistema muestra la pantalla de confirmación de compra y el usuario cierra sesión (logout)")
    elif title == "Multiple Scenarios - 1 Feature":
        doc.add_paragraph(f"""Esta funcionalidad cubre múltiples escenarios de compra en la aplicación para un usuario autenticado:
        1.	Compra de 4 productos: El usuario selecciona 4 productos, ingresa sus datos de envío y pago, y finaliza la compra.
        2.	Compra de 1 producto: Similar al escenario anterior, pero con solo 1 producto.
        3.	Borrar todos los productos del carrito: El usuario agrega 3 productos al carrito, los visualiza y luego los elimina.
        4.	Compra fallida por datos de tarjeta inválidos: El usuario intenta comprar 1 producto, pero deja vacío el número de tarjeta, lo que genera un error de validación.

        Cada escenario finaliza con el cierre de sesión (logout).""")

    for screenshot in screenshots:
        doc.add_paragraph(screenshot)  # Add filename as text
        try:
            doc.add_picture(screenshot, width=Inches(2))  # Add image if exists
        except:
            doc.add_paragraph("[Image Not Found]")  # Handle missing images

# Save the Word file
doc.save("test_report.docx")
print("Word document created: test_report.docx")